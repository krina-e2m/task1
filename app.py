"""
DOCX Generation Server — Flask version for cloud deployment (Render/Railway)
POST /generate  →  { paragraphs: [...] }  →  { success: true, docx_base64: "..." }
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import base64
import traceback
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app)

# ── Styling helpers ───────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def style_run(run, bold, italic, font_size, color_hex=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)

def add_paragraph(doc, text, style, bold, italic, font_size):
    style_map = {
        "Heading 1": "Heading 1", "Heading 2": "Heading 2",
        "Heading 3": "Heading 3", "Title": "Title",
        "Normal": "Normal", "List Bullet": "List Bullet",
        "List Number": "List Number",
    }
    safe = style_map.get(style, "Normal")
    try:
        para = doc.add_paragraph(style=safe)
    except Exception:
        para = doc.add_paragraph(style="Normal")
    if text:
        run = para.add_run(text)
        style_run(run, bold, italic, font_size)
    return para

def build_table(doc, header_item, row_items):
    header_cols = [c.strip() for c in header_item["text"].split(" | ")]
    num_cols = len(header_cols)

    parsed_rows = []
    for item in row_items:
        cols = [c.strip() for c in item["text"].split(" | ")]
        while len(cols) < num_cols:
            cols.append("N/A")
        parsed_rows.append({"cols": cols[:num_cols], "italic": item.get("italic", False)})

    table = doc.add_table(rows=1 + len(parsed_rows), cols=num_cols)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, col_text in enumerate(header_cols):
        cell = hdr_row.cells[i]
        cell.text = ""
        set_cell_bg(cell, "1F3864")
        set_cell_border(cell)
        run = cell.paragraphs[0].add_run(col_text)
        style_run(run, bold=True, italic=False, font_size=11, color_hex="FFFFFF")

    # Data rows
    for row_idx, row_data in enumerate(parsed_rows):
        tr = table.rows[row_idx + 1]
        bg = "EEF2F7" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, col_text in enumerate(row_data["cols"]):
            cell = tr.cells[col_idx]
            cell.text = ""
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            run = cell.paragraphs[0].add_run(col_text)
            style_run(run, bold=False, italic=row_data["italic"], font_size=11, color_hex="000000")

    doc.add_paragraph()

def build_docx(paragraphs):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    i = 0
    while i < len(paragraphs):
        item = paragraphs[i]
        style = item.get("style", "Normal")

        if style == "Table Header":
            row_items = []
            j = i + 1
            while j < len(paragraphs) and paragraphs[j].get("style") == "Table Row":
                row_items.append(paragraphs[j])
                j += 1
            build_table(doc, item, row_items)
            i = j
            continue

        if style == "Table Row":
            add_paragraph(doc, item.get("text", ""), "Normal",
                         item.get("bold", False), item.get("italic", False),
                         item.get("font_size", 11))
            i += 1
            continue

        add_paragraph(doc, item.get("text", "").strip(), style,
                     item.get("bold", False), item.get("italic", False),
                     item.get("font_size", 11))
        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Routes ────────────────────────────────────────────────────────────

@app.route("/", methods=["GET"])
def health():
    return jsonify({"status": "DOCX server is running"}), 200

@app.route("/generate", methods=["POST"])
def generate():
    try:
        data = request.get_json(force=True)
        paragraphs = data if isinstance(data, list) else data.get("paragraphs", [])
        if not paragraphs:
            return jsonify({"success": False, "error": "No paragraphs received"}), 400
        docx_bytes = build_docx(paragraphs)
        b64 = base64.b64encode(docx_bytes).decode()
        return jsonify({"success": True, "docx_base64": b64}), 200
    except Exception as e:
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5050, debug=False)
